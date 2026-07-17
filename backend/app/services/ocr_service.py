# ─────────────────────────────────────────────────────────────
#  app/services/ocr_service.py  —  Text extraction (Hruthi)
# ─────────────────────────────────────────────────────────────
import os
from loguru import logger

try:
    import pdfplumber
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False
    logger.warning("pdfplumber not installed — PDF text extraction disabled")

try:
    from PIL import Image
    import pytesseract
    # On Windows, point to the Tesseract executable
    if os.name == "nt":
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
    logger.warning("pytesseract / Pillow not installed — image OCR disabled")

try:
    import docx as _docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx not installed — DOCX text extraction disabled")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF using pdfplumber (native text layer)."""
    if not _PDF_AVAILABLE:
        return ""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_image(file_path: str) -> str:
    """Extract text from an image using Tesseract OCR."""
    if not _OCR_AVAILABLE:
        return ""
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word (.docx) document."""
    if not _DOCX_AVAILABLE:
        return ""
    document = _docx.Document(file_path)
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """
    Extract text from any supported file.
    - PDFs: try native text layer first; fall back to page-image OCR for scanned docs.
    - Images (JPEG, PNG, TIFF): direct Tesseract OCR.
    - Word documents (.docx): paragraph + table text via python-docx.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        # Scanned PDF has no text layer — convert pages to images and OCR
        if not text.strip() and _OCR_AVAILABLE and _PDF_AVAILABLE:
            logger.info(f"Scanned PDF detected, switching to image OCR: {file_path}")
            try:
                with pdfplumber.open(file_path) as pdf:
                    ocr_parts = []
                    for page in pdf.pages:
                        img = page.to_image(resolution=300).original
                        ocr_parts.append(pytesseract.image_to_string(img))
                    text = "\n".join(ocr_parts)
            except Exception as e:
                logger.error(f"Scanned PDF OCR failed: {e}")
        return text

    elif ext in {".jpg", ".jpeg", ".png", ".tiff", ".tif"}:
        return extract_text_from_image(file_path)

    elif ext == ".docx":
        return extract_text_from_docx(file_path)

    else:
        logger.warning(f"Unsupported file type for OCR: {ext}")
        return ""
